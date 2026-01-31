from docx import Document
from pathlib import Path
from datetime import datetime
from docx.shared import Pt
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE_DIR / "templates" / "service_task_template.docx"
OUTPUT_DIR = BASE_DIR / "generated"

OUTPUT_DIR.mkdir(exist_ok=True)


def set_paragraph_font(paragraph, size=9):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def replace_placeholders(paragraph, data: dict):
    text = paragraph.text

    replaced = False
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in text:
            text = text.replace(placeholder, str(value))
            replaced = True

    if replaced:
        paragraph.clear()
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def generate_service_task(data: dict) -> Path:
    doc = Document(TEMPLATE_PATH)

    # ─── обычный текст ───
    for paragraph in doc.paragraphs:
        replace_placeholders(paragraph, data)

    # ─── таблицы ───
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders(paragraph, data)

    filename = (
        f"service_task_"
        f"{data['city']}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )

    output_path = OUTPUT_DIR / filename
    doc.save(output_path)

    return output_path
