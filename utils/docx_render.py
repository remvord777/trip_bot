from docx import Document
from pathlib import Path
from datetime import datetime


BASE_DIR = Path(__file__).resolve().parents[1]
TEMPLATES_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "generated"

OUTPUT_DIR.mkdir(exist_ok=True)


FILE_TITLES = {
    "service_task.docx": "служебное_задание",
    "money_avans.docx": "заявление_командировка",
    "advance_report.docx": "авансовый_отчет",   # ← НОВОЕ
}


def _replace_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs)
    replaced = False

    for key, value in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, value)
            replaced = True

    if not replaced:
        return

    # очищаем runs
    for run in paragraph.runs:
        run.text = ""

    # записываем в первый run
    paragraph.runs[0].text = full_text



def render_docx(template_name: str, data: dict) -> Path:
    template_path = TEMPLATES_DIR / template_name
    doc = Document(template_path)

    replacements = {
        "{{employee_name}}": data.get("employee_name", ""),
        "{{employee_short}}": data.get("employee_short", ""),
        "{{position}}": data.get("position", ""),
        "{{city}}": f"{data.get('settlement_prefix', '')} {data.get('city', '')}".strip(),
        "{{object}}": data.get("object_name", ""),
        "{{contract}}": data.get("contract", ""),
        "{{date_from}}": data.get("date_from", ""),
        "{{date_to}}": data.get("date_to", ""),
        "{{total}}": str(data.get("total", "")),
        "{{purpose}}": data.get("service", ""),
        "{{advance_amount}}": str(data.get("advance_amount", "")),
        "{{apply_date}}": datetime.now().strftime("%d.%m.%Y"),


        # --- НОВОЕ ДЛЯ АВАНСОВОГО ---
        "{{report_date}}": datetime.now().strftime("%d.%m.%Y"),
        "{{department}}": data.get("department", ""),
        "{{object_name}}": data.get("object_name", ""),

        "{{per_diem_text}}": (
            f"Суточные {data.get('per_diem_rate', 0)} × {data.get('total', 0)} дней"
        ),
        "{{per_diem_total}}": f"{data.get('per_diem_total', 0):,}".replace(",", " "),

        "{{accommodation_amount}}": f"{data.get('accommodation_amount', 0):,}".replace(",", " "),
        "{{taxi_amount}}": f"{data.get('taxi_amount', 0):,}".replace(",", " "),
        "{{ticket_amount}}": f"{data.get('ticket_amount', 0):,}".replace(",", " "),

        "{{total_amount}}": f"{data.get('total_amount', 0):,}".replace(",", " "),
    }



    # ===== ПАРАГРАФЫ =====
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # ===== ТАБЛИЦЫ =====
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # ===== ИМЯ ФАЙЛА =====
    doc_type = FILE_TITLES.get(template_name, "документ")

    employee_name = data.get("employee_name", "").strip()
    surname = employee_name.split()[0] if employee_name else "БезФИО"

    date_from = data.get("date_from", "")
    date_to = data.get("date_to", "")

    filename = f"ИМ_{doc_type}_{surname}_{date_from}–{date_to}.docx"

    output_path = OUTPUT_DIR / filename
    doc.save(output_path)

    return output_path

