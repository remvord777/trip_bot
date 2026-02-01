from docx import Document
from datetime import date
from pathlib import Path

TEMPLATE_PATH = Path("templates/money_avans.docx")
OUTPUT_DIR = Path("generated")
OUTPUT_DIR.mkdir(exist_ok=True)


def _replace_in_paragraphs(paragraphs, replacements: dict):
    for paragraph in paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))


def _replace_in_tables(tables, replacements: dict):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, replacements)
                _replace_in_tables(cell.tables, replacements)  # на будущее


def generate_advance_request(data: dict) -> Path:
    doc = Document(TEMPLATE_PATH)

    today = date.today().strftime("%d.%m.%Y")

    replacements = {
        "{{employee_name}}": data["employee_name"],
        "{{apply_date}}": today,
        "{{city}}": data["city"],
        "{{object}}": data["object"],
        "{{contract}}": data["contract"],
        "{{date_from}}": data["date_from"],
        "{{date_to}}": data["date_to"],
        "{{advance_amount}}": data["advance_amount"],
    }

    # 🔥 КЛЮЧЕВОЕ
    _replace_in_paragraphs(doc.paragraphs, replacements)
    _replace_in_tables(doc.tables, replacements)

    filename = f"advance_{today}_{data['employee_name']}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(output_path)

    return output_path
